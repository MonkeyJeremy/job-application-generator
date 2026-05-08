"""
Renders a structured resume dict into a .docx matching Jeremy's reference format:
- Times New Roman throughout
- 0.40" L/R margins, 0.457" T/B margins
- Justified alignment
- No explicit paragraph spacing (Word defaults = 0 before/after, single line)
- Hanging indent 0.158" for bullets and skills
- Bottom border (sz=4) on section headers
"""
import io
import os
import tempfile

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from resume_context import CANDIDATE_NAME, CANDIDATE_EMAIL, CANDIDATE_PHONE

FONT = "Times New Roman"
CONTENT_WIDTH_TWIPS = 11088  # 8.5" - 0.4" - 0.4" = 7.7" × 1440


def _add_run(para, text, bold=False, italic=False, size=10):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return r


def _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    # No explicit spacing — inherits from Normal (0/0/single)
    return p


def _add_border(para):
    """Adds bottom border to a paragraph (matching reference sz=4)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc, title):
    p = _para(doc)
    _add_run(p, title, bold=True, size=11)
    _add_border(p)
    return p


def _two_col_row(doc, left, right, left_bold=False, left_italic=False):
    """Company/location or title/dates on one line using a right tab stop."""
    p = _para(doc)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(CONTENT_WIDTH_TWIPS))
    tabs.append(tab)
    pPr.append(tabs)
    _add_run(p, left, bold=left_bold, italic=left_italic)
    _add_run(p, "\t")
    _add_run(p, right, italic=left_italic)
    return p


def _bullet(doc, text):
    p = _para(doc)
    p.paragraph_format.left_indent = Inches(0.158)
    p.paragraph_format.first_line_indent = Inches(-0.158)
    _add_run(p, text)
    return p


def _skill_line(doc, category, items):
    p = _para(doc)
    p.paragraph_format.left_indent = Inches(0.158)
    p.paragraph_format.first_line_indent = Inches(-0.158)
    _add_run(p, f"{category}: ", bold=True)
    _add_run(p, items)
    return p


def resume_to_docx(data: dict) -> bytes:
    doc = Document()

    # Margins matching reference
    for section in doc.sections:
        section.top_margin = Inches(0.457)
        section.bottom_margin = Inches(0.457)
        section.left_margin = Inches(0.400)
        section.right_margin = Inches(0.400)

    # Normal style: Times New Roman 10pt, 0 spacing, single line
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ── Name ──────────────────────────────────────────────────────────────
    name_p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    name_p.paragraph_format.space_before = Pt(5)
    r = name_p.add_run(CANDIDATE_NAME)
    r.font.name = FONT
    r.font.size = Pt(14)
    r.font.bold = True
    # line spacing 1.075× to match reference
    pPr = name_p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(1.075 * 240)))  # 240 twips = single
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    # ── Contact ───────────────────────────────────────────────────────────
    contact_p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(contact_p, f"{CANDIDATE_PHONE} | {CANDIDATE_EMAIL} | U.S. Citizen")

    # ── Education ─────────────────────────────────────────────────────────
    _section_header(doc, "EDUCATION")
    _two_col_row(doc, "University of California, San Diego", "La Jolla, CA", left_bold=True)
    _two_col_row(doc, "Bachelor of Science in Data Science, Business Minor",
                 "Sept 2018 – Sept 2024", left_italic=True)
    p = _para(doc)
    p.paragraph_format.left_indent = Inches(0.158)
    p.paragraph_format.first_line_indent = Inches(-0.158)
    _add_run(p,
             "Relevant Coursework: Data Analytics, Business Analytics, Strategic Planning, "
             "Financial Analytics, Risk Assessment, Data Structures & Algorithms, Probability & "
             "Statistics, Machine Learning & Deep Learning, Data Science, Market Management",
             italic=True)

    # ── Skills ────────────────────────────────────────────────────────────
    _section_header(doc, "SKILLS")
    for category, items in data.get("skills", {}).items():
        _skill_line(doc, category, items)

    # ── Summary (optional) ────────────────────────────────────────────────
    summary = data.get("summary", "").strip()
    if summary:
        _section_header(doc, "SUMMARY")
        _bullet(doc, summary)

    # ── Professional Experience ───────────────────────────────────────────
    _section_header(doc, "PROFESSIONAL EXPERIENCE")
    for job in data.get("experience", []):
        _two_col_row(doc, job["company"], job["location"], left_bold=True)
        _two_col_row(doc, job["title"], job["dates"], left_italic=True)
        for b in job.get("bullets", []):
            _bullet(doc, b)

    # ── Project Experience ────────────────────────────────────────────────
    _section_header(doc, "PROJECT EXPERIENCE")
    for proj in data.get("projects", []):
        _two_col_row(doc, proj["name"], proj.get("dates", ""), left_bold=True)
        for b in proj.get("bullets", []):
            _bullet(doc, b)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def resume_to_pdf(docx_bytes: bytes) -> bytes | None:
    try:
        from docx2pdf import convert
    except ImportError:
        return None

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "resume.docx")
        pdf_path = os.path.join(tmp, "resume.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            return None
