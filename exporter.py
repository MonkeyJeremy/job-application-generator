import io
import os
import tempfile
from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from resume_context import CANDIDATE_NAME, CANDIDATE_EMAIL, CANDIDATE_PHONE


def _p(doc, text, size=11, bold=False, space_after=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    return para


def to_docx(letter_body: str, company: str = "") -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Header ────────────────────────────────────────────────────────────
    _p(doc, CANDIDATE_NAME, size=14, bold=True, space_after=0,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, f"{CANDIDATE_PHONE} | {CANDIDATE_EMAIL}", size=11, space_after=18,
       align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Date ──────────────────────────────────────────────────────────────
    today = date.today().strftime("%B %#d, %Y") if os.name == "nt" else date.today().strftime("%B %-d, %Y")
    _p(doc, today, space_after=12)

    # ── Recipient ─────────────────────────────────────────────────────────
    if company:
        _p(doc, "Hiring Team", space_after=0)
        _p(doc, company, space_after=18)

    # ── Body ──────────────────────────────────────────────────────────────
    # Split off the closing block ("Sincerely," onward) from the body
    closing_marker = None
    for marker in ("Sincerely,", "Best regards,", "Best,", "Regards,"):
        if marker in letter_body:
            closing_marker = marker
            break

    if closing_marker:
        body_part, closing_part = letter_body.split(closing_marker, 1)
    else:
        body_part = letter_body
        closing_part = ""

    body_paras = [p.strip() for p in body_part.split("\n\n") if p.strip()]
    for i, text in enumerate(body_paras):
        # Last body paragraph gets extra space before the closing
        space = 18 if i == len(body_paras) - 1 else 12
        _p(doc, text, space_after=space)

    # ── Closing ───────────────────────────────────────────────────────────
    _p(doc, closing_marker or "Sincerely,", space_after=0)
    _p(doc, "", space_after=12)   # blank line between closing and name
    _p(doc, CANDIDATE_NAME, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(docx_bytes: bytes) -> bytes | None:
    try:
        from docx2pdf import convert
    except ImportError:
        return None

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "cover_letter.docx")
        pdf_path = os.path.join(tmp, "cover_letter.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            return None
